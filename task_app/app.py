import streamlit as st
from docx import Document
from io import BytesIO
import smtplib
from email.message import EmailMessage

st.set_page_config(page_title="Оценка задачи", layout="wide")

st.title("📊 Оценка задачи")

# --- 0. НАЗВАНИЕ ---
project_name = st.text_input("0. Название проекта / задачи *")

# --- 1. ТИП ---
task_type = st.selectbox("1. Что вы хотите получить?", [
    "Дашборд",
    "Отчёт / таблица",
    "Показатель (метрика)",
    "Таблица KPI",
    "Алерт / рассылка",
    "Поиск дублей / сверка",
    "Справочник (MDM)",
    "Интеграция",
    "Другое"
])

# --- 2. ИСТОЧНИКИ ---
sources = st.multiselect("2. Откуда будут браться данные?", [
    "1С", "CRM", "MDM", "Directum", "API",
    "Веб-интерфейс", "Веб-сервис", "Несколько источников", "Пока неизвестно"
])

# --- 3. ПОДГОТОВКА ---
data_prep = st.selectbox("3. Нужно ли подготавливать данные?", [
    "Нет, данные готовы",
    "Нужно немного доработать",
    "Нужно сильно перерабатывать / объединять"
])

# --- 4. ОБЪЕМ ---
volume = st.selectbox("4. Какой объём данных?", [
    "Небольшой",
    "Средний",
    "Большой"
])

# --- 5. СЛОЖНОСТЬ ---
complexity = st.selectbox("5. Насколько сложная логика?", [
    "Простая",
    "Средняя",
    "Сложная"
])

# --- 6. ДОСТУП ---
access = st.selectbox("6. Нужно ли ограничение доступа?", [
    "Нет",
    "Да"
])

# --- 7. ОБНОВЛЕНИЕ ---
refresh = st.selectbox("7. Как часто обновляются данные?", [
    "Разово",
    "Раз в день",
    "В течение дня"
])

# --- 8. ТРЕБОВАНИЯ ---
requirements = st.selectbox("8. Есть ли описание требований?", [
    "Нет",
    "Есть частично",
    "Есть подробное описание"
])

req_files = st.file_uploader("📎 Приложите файлы требований", accept_multiple_files=True)

# --- 9. ПРОТОТИП ---
prototype = None
proto_files = None

if task_type == "Дашборд":
    prototype = st.selectbox("9. Есть ли прототип дашборда?", [
        "Нет",
        "Есть пример",
        "Есть подробный прототип"
    ])

    proto_files = st.file_uploader("📎 Прототип", accept_multiple_files=True)

# --- 10. ПРИОРИТЕТ ---
st.subheader("10. Оценка приоритета (0–2)")

st.markdown("""
**1. Влияние на бизнес**  
0 — Локально  
1 — Средне  
2 — KPI / выручка  

**2. Формулировка**  
0 — Нет ТЗ  
1 — Частично  
2 — Чёткое ТЗ  

**3. Реализуемость**  
0 — R&D  
1 — Нужно изучение  
2 — Всё готово  

**4. Риски**  
0 — Высокие  
1 — Средние  
2 — Низкие  

**5. Стоимость**  
0 — >100 ч  
1 — 50–100 ч  
2 — <50 ч  
""")

impact = st.slider("Влияние на бизнес", 0, 2)
clarity = st.slider("Формулировка", 0, 2)
realization = st.slider("Реализуемость", 0, 2)
risks = st.slider("Риски", 0, 2)
cost = st.slider("Стоимость", 0, 2)

priority = impact + clarity + realization + risks + cost

# --- РАСЧЕТ ЧАСОВ ---
def calculate_hours():
    estimate = 2
    analysis = 20
    arch = 4
    dev = 40
    test = 16
    qa = 8

    if "немного" in data_prep:
        analysis *= 1.1
    if "сильно" in data_prep:
        analysis *= 1.3
        dev *= 1.3

    if complexity == "Средняя":
        dev *= 1.2
    if complexity == "Сложная":
        dev *= 1.5

    if requirements == "Есть подробное описание":
        analysis *= 1.3

    if task_type == "Дашборд" and prototype == "Есть подробный прототип":
        dev *= 0.9

    if len(sources) > 1 or "Несколько источников" in sources:
        analysis *= 1.2

    if complexity == "Сложная":
        dev *= 1.2
        test *= 1.2

    return {
        "estimate": round(estimate),
        "analysis": round(analysis),
        "arch": round(arch),
        "dev": round(dev),
        "test": round(test),
        "qa": round(qa),
        "total": round(estimate + analysis + arch + dev + test + qa)
    }

# --- СОЗДАНИЕ DOC ---
def create_doc(hours):
    doc = Document()

    doc.add_heading(project_name, 0)

    doc.add_heading("Описание", 1)
    doc.add_paragraph(f"Тип: {task_type}")
    doc.add_paragraph(f"Источники: {', '.join(sources)}")
    doc.add_paragraph(f"Подготовка данных: {data_prep}")
    doc.add_paragraph(f"Объем: {volume}")
    doc.add_paragraph(f"Сложность: {complexity}")
    doc.add_paragraph(f"Обновление: {refresh}")
    doc.add_paragraph(f"Требования: {requirements}")

    if task_type == "Дашборд":
        doc.add_paragraph(f"Прототип: {prototype}")

    doc.add_heading("Приоритет", 1)
    doc.add_paragraph(f"""
Влияние: {impact}
Формулировка: {clarity}
Реализуемость: {realization}
Риски: {risks}
Стоимость: {cost}

ИТОГО: {priority}
""")

    doc.add_heading("Оценка трудозатрат", 1)
    doc.add_paragraph(f"""
Оценка: {hours['estimate']}
Анализ: {hours['analysis']}
Архитектура: {hours['arch']}
Разработка: {hours['dev']}
Тестирование: {hours['test']}
Выверка: {hours['qa']}

ИТОГО: {hours['total']}
""")

    doc.add_heading("Приложения", 1)

    if req_files:
        for f in req_files:
            doc.add_paragraph(f"Требования: {f.name}")

    if proto_files:
        for f in proto_files:
            doc.add_paragraph(f"Прототип: {f.name}")

    file = BytesIO()
    doc.save(file)
    file.seek(0)
    return file

# --- ОТПРАВКА EMAIL ---
def send_email(doc_file, hours):
    msg = EmailMessage()
    msg["Subject"] = f"Заявка: {project_name}"
    msg["From"] = "your_email@company.ru"
    msg["To"] = "antonavstrijchenko@iek.ru"

    msg.set_content(f"""
Название: {project_name}
Приоритет: {priority}
Оценка: {hours['total']} часов
""")

    msg.add_attachment(
        doc_file.getvalue(),
        maintype="application",
        subtype="vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=f"{project_name}.docx"
    )

    with smtplib.SMTP("smtp.your_company.ru", 25) as server:
        server.send_message(msg)

# --- КНОПКА ---
if st.button("📊 Сформировать заявку"):

    if not project_name:
        st.error("Введите название проекта")
        st.stop()

    hours = calculate_hours()

    st.success("Готово")

    st.write(f"### {project_name}")

    st.write("### 📊 Приоритет")
    st.write(f"""
Влияние: {impact}
Формулировка: {clarity}
Реализуемость: {realization}
Риски: {risks}
Стоимость: {cost}

ИТОГО: {priority}
""")

    st.write("### ⏱ Оценка")
    st.write(f"""
Оценка: {hours['estimate']}
Анализ: {hours['analysis']}
Архитектура: {hours['arch']}
Разработка: {hours['dev']}
Тестирование: {hours['test']}
Выверка: {hours['qa']}

ИТОГО: {hours['total']}
""")

    doc_file = create_doc(hours)

    st.download_button(
        "📄 Скачать Word",
        data=doc_file,
        file_name=f"{project_name}.docx"
    )

    # отправка (включай когда настроишь SMTP)
    # send_email(doc_file, hours)